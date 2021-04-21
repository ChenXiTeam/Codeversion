# -*- coding: utf-8 -*-
import docx
from docx import Document
import base64
import requests
import re
import os
from os.path import basename
import chardet
import fuzzywuzzy
from fuzzywuzzy import fuzz

# 返回不含后缀和路径的文件名
def get_file_name(s):
    index_head = s.rfind("/")
    index_hail = s.rfind(".")
    result = s[index_head + 1:index_hail]
    return result


# 返回不含后缀的文件名
def get_no_tail_file_name(s):
    index_hail = s.rfind(".")
    result = s[:index_hail]
    return result


# 提取docx
# 将docx转为txt
def for_docx(in_path, out_path):
    ft = open(out_path, 'w', encoding='utf-8')  #全部内容
    document = Document(in_path)

    all_text = ""
    # 段落
    for paragraph in document.paragraphs:
        all_text = all_text + paragraph.text

    # 表格
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]  # 获取文件中的第一个表格
    for i in range(0, len(table.rows)):  # 从表格第一行开始循环读取表格数据
        result = f'{(table.cell(i, 0).text):<5}'
        # cell(i,0)表示第(i+1)行第1列数据,以此类推
        all_text = all_text + result

    ft.write(all_text)

    ft.close()

    type = get_marked_code(all_text, out_path)

    return type


# 获取匹配到的代码
def get_marked_code(all_text, out_path):
    # 正则表达式处理
    # 方法比较蠢，伺机优化
    # 提取有标记的代码

    r1 = r'--begin--c--code--(.+?)(?=--end--c--code--)'
    r2 = r'--begin--cpp--code--(.+?)(?=--end--cpp--code--)'
    r3 = r'--begin--java--code--(.+?)(?=--end--java--code--)'
    r4 = r'--begin--python--code--(.+?)(?=--end--python--code--)'

    pattern1 = re.compile(r1, re.DOTALL)
    pattern2 = re.compile(r2, re.DOTALL)
    pattern3 = re.compile(r3, re.DOTALL)
    pattern4 = re.compile(r4, re.DOTALL)

    procedure1 = pattern1.findall(all_text)
    procedure2 = pattern2.findall(all_text)
    procedure3 = pattern3.findall(all_text)
    procedure4 = pattern4.findall(all_text)

    t = 1
    if (len(procedure1) != 0):
        print('1')
        for code in procedure1:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.c', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()
            print(t)
            make_to_string(out_path[:-4] + '_marked_code_' + t.__str__() + '.c',
                           out_path[:-4] + '_big_string_' + t.__str__() + '.txt')
        return 2

    if (len(procedure2) != 0):
        print('2')
        for code in procedure2:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.cpp', 'w', encoding='utf-8')
            fcf.write(code)
            fcf.close()
            make_to_string(out_path[:-4] + '_marked_code_' + t.__str__() + '.cpp',
                           out_path[:-4] + '_big_string_' + t.__str__() + '.txt')
            t = t + 1
        return 2

    if (len(procedure3) != 0):
        print('3')
        for code in procedure3:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.java', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()
            make_to_string(out_path[:-4] + '_marked_code_' + t.__str__() + '.java',
                           out_path[:-4] + '_big_string_' + t.__str__() + '.txt')
        return 1

    if (len(procedure4) != 0):
        print('4')
        for code in procedure4:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.python', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()
            make_to_string(out_path[:-4] + '_marked_code_' + t.__str__() + '.python',
                           out_path[:-4] + '_big_string_' + t.__str__() + '.txt')
        return 3



# 提取docx中的图片
def get_pictures(word_path, result_path):
    """
    图片提取
    :param word_path: word路径(---.docx)
    :param result_path: 结果路径（---.txt）
    :return:
    """
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            if not os.path.exists(result_path):
                os.makedirs(result_path)
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = f'{new_name}_{img_name}'
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)



# 提取图片内容
def for_picture(in_path, out_path):
    url = 'https://aip.baidubce.com/oauth/2.0/token'
    data = {
        'grant_type': 'client_credentials',  # 固定值
        'client_id': '3BKl4cKDe9wHIivL5TMirQBd',  # 在开放平台注册后所建应用的API Key
        'client_secret': 'sgYVAVXdtv2TyhqwCuRkYur0P4tWlHyc'  # 所建应用的Secret Key
    }
    res = requests.post(url, data=data)
    res = res.json()
    access_token = res['access_token']

    # 通用文字识别接口url
    general_word_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"

    # 二进制方式打开图片文件
    f = open(in_path, 'rb')
    img = base64.b64encode(f.read())
    params = {"image": img,
              "language_type": "CHN_ENG"}
    request_url = general_word_url + "?access_token=" + access_token
    headers = {'content-type': 'application/x-www-form-urlencoded'}
    response = requests.post(request_url, data=params, headers=headers)

    if response:
        res = response.json()["words_result"]
        with open(out_path, 'w', encoding='utf-8') as f:
            for j in res:
                f.write(j["words"] + "\n")


# 判断txt中的代码类型,想法没用到
def txt_to_code(in_path, out_path):
    # 结合图片识别使用（文本识别已经可以生成源代码文件）
    # 由于一小块代码意义不大，这里暂时考虑完整的代码
    with open(in_path, encoding="utf-8") as file:
        alltext = file.read()
        print(alltext)

        # 指针回到开头
        file.seek(0)

        # 读文本的前10个字符并打印出来
        head = file.read(10)
        print(head)

        # 部分匹配，如果S1是S2的子串依然返回100
        # 没有区分c和c++
        same_degree = []

        same_degree.append(fuzz.partial_ratio(head, '#include <iostream>'))
        same_degree.append(fuzz.partial_ratio(head, 'using namespace std;'))
        same_degree.append(fuzz.partial_ratio(head, 'package'))
        same_degree.append(fuzz.partial_ratio(head, 'public static void main'))
        same_degree.append(fuzz.partial_ratio(head, 'public void'))
        same_degree.append(fuzz.partial_ratio(head, 'int main()'))
        same_degree.append(fuzz.partial_ratio(head, 'def'))

        max_data = max(same_degree)
        max_index = same_degree.index(max(same_degree))+1

        tail = ''
        if max_index in range(1,3):
            tail = 'cpp'
        if max_index in range(3, 6):
            tail = 'java'
        if max_index in range(6, 8):
            tail = 'python'

        out_path = in_path[:-3]+tail

        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(alltext)
        f.close()
    file.close()



# 获取文件编码格式
def get_encode(path):
  f = open(path, 'rb')
  data = f.read()
  f.close()
  encode = (chardet.detect(data))['encoding']
  return encode


# 去掉代码中的注释和空行
def make_to_string(in_path, out_path):

  if in_path.split('.')[-1] == 'py':
      bds = '#.*'
      target = re.compile(bds)  # 单行注释
      encode = get_encode(in_path)

      if (encode == 'utf-8'):
          f = open(in_path, encoding='utf-8')
      else:
          f = open(in_path)

      data = f.read()

      result0 = target.findall(data)

      result = []
      result += result0
      for i in result:
          data = data.replace(i, '')  # 替换为空字符串

      st = list(data)


  elif in_path.split('.')[-1] == 'c' or in_path.split('.')[-1] == 'cpp' or in_path.split('.')[-1] == 'java':
      bds0 = '//.*'  # 标准匹配单行注释
      bds1 = '\/\*(?:[^\*]|\*+[^\/\*])*\*+\/'  # 标准匹配多行注释  可匹配跨行注释
      target0 = re.compile(bds0)  # 单行注释
      target = re.compile(bds1)  # 编译正则表达式
      encode = get_encode(in_path)

      if (encode == 'utf-8'):
          f = open(in_path, encoding='utf-8')
      else:
          f = open(in_path)

      data = f.read()

      result0 = target0.findall(data)

      result = target.findall(data)

      result += result0
      for i in result:
          data = data.replace(i, '')  # 替换为空字符串

      st = list(data)


  # 去掉空格一行换行
  for i in range(0, len(st)):
    for line in st:
      if '\n' in line:
        index = st.index(line)
        del st[index]
      if ' ' in line:
        index = st.index(line)
        del st[index]
      if '\t' in line:
        index = st.index(line)
        del st[index]
    mn = "".join(st)

  file = open(out_path, 'w', encoding='utf-8')
  file.write(mn)
  file.close()


# 去掉代码中的变量名
def delete_var(in_path, out_path):
    print()








