import docx
from docx import Document

import re

in_path = "E:\python\py_pick\\test\\test2.docx"
out_path_text = 'E:\python\py_pick\\result\\test2text.txt'
out_path_code = 'E:\python\py_pick\\result\\test2code.txt'

ft = open(out_path_text, 'w' ,encoding='utf-8')
fc = open(out_path_code, 'w' ,encoding='utf-8')

document = Document(in_path)

all_text = ""

#段落
for paragraph in document.paragraphs:
    all_text = all_text + paragraph.text

#表格
tables = document.tables   # 获取文件中的表格集
table = tables[0]  # 获取文件中的第一个表格
for i in range(0, len(table.rows)):  # 从表格第一行开始循环读取表格数据
    result = f'{(table.cell(i, 0).text):<5}'
    # cell(i,0)表示第(i+1)行第1列数据,以此类推
    all_text = all_text + result


#正则表达式处理
pattern = re.compile(r'--begin--code--(.+?)(?=--end--code--)', re.DOTALL)
procedure = pattern.findall(all_text)

for code in procedure:
    fc.write(code)

ft.write(all_text)

ft.close()
fc.close()