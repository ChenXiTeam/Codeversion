import docx
from docx import Document

in_path = "E:\\test\\test1.docx"
out_path = 'E:\\test\\test1.txt'

f = open(out_path, 'w' ,encoding='utf-8')
document = Document(in_path)

#段落
for paragraph in document.paragraphs:
    print(paragraph.text)
    f.write(paragraph.text)

#表格
tables = document.tables   # 获取文件中的表格集
table = tables[0]  # 获取文件中的第一个表格
for i in range(0, len(table.rows)):  # 从表格第一行开始循环读取表格数据
    result = f'{(table.cell(i, 0).text):<5}' + "" + f'{(table.cell(i, 1).text):<5}' + "" + f'{(table.cell(i, 2).text):<5}'
    # cell(i,0)表示第(i+1)行第1列数据,以此类推
    print(result)
    f.write(result)

f.close()
