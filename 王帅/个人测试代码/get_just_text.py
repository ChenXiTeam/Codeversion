import docx
from docx import Document
import re


# 提取docx
# 将docx转为txt
def for_docx(in_path, out_path):
    ft = open(out_path, 'w', encoding='utf-8')  #全部内容
    document = Document(in_path)

    all_text = ""
    # 段落
    for paragraph in document.paragraphs:
        all_text = all_text + paragraph.text

    # 表格
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]  # 获取文件中的第一个表格
    for i in range(0, len(table.rows)):  # 从表格第一行开始循环读取表格数据
        result = f'{(table.cell(i, 0).text):<5}'
        # cell(i,0)表示第(i+1)行第1列数据,以此类推
        all_text = all_text + result

    ft.write(all_text)

    ft.close()

    get_marked_code(all_text, out_path)


def get_marked_code(all_text, out_path):
    # 正则表达式处理
    # 方法比较蠢，伺机优化
    # 提取有标记的代码

    r1 = r'--begin--c--code--(.+?)(?=--end--c--code--)'
    r2 = r'--begin--cpp--code--(.+?)(?=--end--cpp--code--)'
    r3 = r'--begin--java--code--(.+?)(?=--end--java--code--)'
    r4 = r'--begin--python--code--(.+?)(?=--end--python--code--)'

    pattern1 = re.compile(r1, re.DOTALL)
    pattern2 = re.compile(r2, re.DOTALL)
    pattern3 = re.compile(r3, re.DOTALL)
    pattern4 = re.compile(r4, re.DOTALL)

    procedure1 = pattern1.findall(all_text)
    procedure2 = pattern2.findall(all_text)
    procedure3 = pattern3.findall(all_text)
    procedure4 = pattern4.findall(all_text)

    t = 1
    if (len(procedure1) != 0):
        print('1')
        for code in procedure1:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.c', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()

    if (len(procedure2) != 0):
        print('2')
        for code in procedure2:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.cpp', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()

    if (len(procedure3) != 0):
        print('3')
        for code in procedure3:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.java', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()

    if (len(procedure4) != 0):
        print('4')
        for code in procedure4:
            # 输出代码文件
            fcf = open(out_path[:-4] + '_marked_code_' + t.__str__() + '.python', 'w', encoding='utf-8')
            t = t + 1
            fcf.write(code)
            fcf.close()



in_path='E:\python\py_pick\\result\\test1.docx'
out_path='E:\python\py_pick\\result\\test1.txt'
for_docx(in_path,out_path)